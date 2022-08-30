from bs4 import BeautifulSoup
import requests
from docx import Document
from PySimpleGUI import PySimpleGUI as sg

sg.theme('Reddit')
layout = [
    [sg.Text('URL'), sg.Input(key='URL')],
    [sg.Button('Baixar')]
]
# url = "https://saikaiscan.com.br/series/king-of-gods-kog?tab=capitulos"
janela = sg.Window('LightDownloader', layout)
while True:
    eventos, valores = janela.read()
    if eventos == sg.WINDOW_CLOSED:
        exit()
    
    if eventos == 'Baixar':
        url = valores['URL']
        break

#headers = {'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_10_1) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/39.0.2171.95 Safari/537.36'}
#input("Digite a url do manga desejado")
homePage = requests.get(url)
soup = BeautifulSoup(homePage.content, 'html.parser')
final_soup = soup.find_all('div', {'class': '__right'})
counter = 0
for tag in final_soup:
    capitulos = tag.find_all("a")
    document = Document()
    for link in capitulos:
        if counter > 1:
            urlCap = link['href']
            titulo = link.find('span', {'class': '__chapters--title'})
            urlText = 'https://saikaiscan.com.br' + urlCap
            document.add_heading(titulo.getText(), 2)
            #entra no capitulo
            lightPage = requests.get(urlText)
            SopaDoTexto = BeautifulSoup(lightPage.content, 'html.parser')
            SopaDoTexto = SopaDoTexto.find_all('p')
            numCap = 0
            for paragrafo in SopaDoTexto:
                texto = paragrafo.getText()
                document.add_paragraph(texto)
                numCap += 1
        counter += 1     
        document.save("Cap" + str(counter) + ".docx")
    